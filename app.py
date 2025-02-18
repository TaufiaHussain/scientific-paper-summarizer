# Install required dependencies (uncomment the next line if needed)
# !pip install streamlit pymupdf pdfplumber transformers wordcloud python-docx

import streamlit as st
import fitz  # PyMuPDF for PDF text extraction
import re
import pandas as pd
import matplotlib.pyplot as plt
from transformers import pipeline
from wordcloud import WordCloud
from docx import Document

# ✅ Debugging: Ensure the Streamlit app is running
st.set_page_config(page_title="Research Paper Summarizer", layout="wide")
st.title("📚 Research Paper Summarizer")
st.write("### ✅ Streamlit App is Running Successfully!")

# Sidebar: Upload PDF
st.sidebar.header("Upload a PDF File")
uploaded_file = st.sidebar.file_uploader("Choose a PDF file", type=["pdf"])

def extract_text_from_pdf(pdf_path, max_pages=15):
    """Extracts specific sections (Research Question, Hypothesis, Results, Conclusion, Discussion)."""
    text = ""
    section_text = {"research_question": "", "hypothesis": "", "key_findings": "", "discussion": "", "conclusion": ""}
    
    with fitz.open(pdf_path) as doc:
        for i, page in enumerate(doc):
            page_text = page.get_text("text")
            text += page_text + "\n"

            if "research question" in page_text.lower():
                section_text["research_question"] += page_text
            elif "hypothesis" in page_text.lower():
                section_text["hypothesis"] += page_text
            elif "results" in page_text.lower() or "findings" in page_text.lower():
                section_text["key_findings"] += page_text
            elif "discussion" in page_text.lower():
                section_text["discussion"] += page_text
            elif "conclusion" in page_text.lower():
                section_text["conclusion"] += page_text

            if i >= max_pages:
                break
    
    return section_text

def summarize_text(text, max_length=250):
    """Summarizes longer text in chunks and merges results."""
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
    if not text.strip():
        return "Not found in the document."
    
    chunks = [text[i:i+2048] for i in range(0, len(text), 2048)]
    summaries = [summarizer(chunk, max_length=max_length, min_length=100, do_sample=False)[0]['summary_text']
                 for chunk in chunks]
    return " ".join(summaries)

def generate_summary(pdf_path, summary_length=250):
    """Processes the PDF and returns a structured summary."""
    sections = extract_text_from_pdf(pdf_path)
    
    structured_summary = {
        "Research Question": summarize_text(sections["research_question"], max_length=summary_length),
        "Hypothesis": summarize_text(sections["hypothesis"], max_length=summary_length),
        "Key Findings": summarize_text(sections["key_findings"], max_length=summary_length),
        "Discussion": summarize_text(sections["discussion"], max_length=summary_length),
        "Conclusion": summarize_text(sections["conclusion"], max_length=summary_length),
    }
    
    return structured_summary

def save_summary_as_word(summary_output, filename="summary.docx"):
    """Saves the structured summary as a Word document."""
    doc = Document()
    doc.add_heading("Scientific Paper Summary", level=1)

    for key, value in summary_output.items():
        doc.add_heading(key, level=2)
        doc.add_paragraph(value)

    doc.save(filename)
    st.success(f"✅ Summary saved as `{filename}`. Click below to download.")
    with open(filename, "rb") as f:
        st.download_button("📥 Download Summary", f, file_name=filename)

# Process PDF and Display Summarized Results
if uploaded_file is not None:
    st.sidebar.success("✅ File uploaded successfully!")
    
    with open("temp.pdf", "wb") as f:
        f.write(uploaded_file.getbuffer())

    with st.spinner("⏳ Processing... Please wait."):
        summary_output = generate_summary("temp.pdf", summary_length=250)

    st.subheader("📌 Structured Summary")
    for key, value in summary_output.items():
        with st.expander(f"🔹 {key}", expanded=True):
            st.write(value)

    # Save as Word document
    save_summary_as_word(summary_output, "summary.docx")

